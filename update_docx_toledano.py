import os, json, glob, shutil
from docx import Document
from docx.shared import Pt

doc_path = r"C:\Users\yejc2\Downloads\Estandar_Respuesta_WebAPI_SAP_ROAD.docx"
payload_dir = r"C:\Users\yejc2\source\repos\ROAD_TOLEDANO\ROADWedAPI\payloads\sap_descuentos"
backup_path = doc_path.replace('.docx', '_backup_20260601.docx')

shutil.copy2(doc_path, backup_path)
doc = Document(doc_path)

doc.add_paragraph('')
doc.add_heading('8. Actualizacion funcional y tecnica (2026-06-01)', level=1)
doc.add_paragraph('Esta seccion integra los cambios implementados en WebAPI para fase PDT Promociones 1.2 (combos y descuentos/recargos SAP-ROAD).')

bullets = [
    'Endpoint operativo adicional habilitado: /api/sap/descuentos.',
    'Se mantiene SapDescuentoController / SapDescuentoMapper como componentes principales de transformacion.',
    'Coddesc (CODDESC ROAD) se recibe en DTO y debe persistirse como identificador de condicion tambien para escenarios simples/escalas cuando venga en payload.',
    'Mapeo confirmado A906/KDGRP -> CTIPO=3 (tipo cliente ROAD).',
    'PTIPO incorpora valor 6 para combos; detalle en P_DESCUENTO_COMBO_DET.',
    'Rama HH de referencia funcional: dev_road_2024_bak3.'
]
for b in bullets:
    p = doc.add_paragraph(b)
    p.style = 'List Bullet'

doc.add_heading('9. Estandar de respuesta recomendado (resumen)', level=1)
resp = '''{
  "success": true,
  "statusCode": "OK",
  "message": "Condicion procesada correctamente",
  "traceId": "ROAD-20260601-000001",
  "timestamp": "2026-06-01T00:00:00Z",
  "result": {
    "coddescRoad": 7101001,
    "coddescSap": "ZK94",
    "regCond": "9101001",
    "promo": "7101001",
    "tipoOperacion": "DESCUENTO",
    "tipoRegistro": "SIMPLE",
    "targetTable": "P_DESCUENTO",
    "recordsAffected": 1
  },
  "warnings": [],
  "errors": []
}'''
pr = doc.add_paragraph(resp)
for run in pr.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

doc.add_heading('10. Payloads por escenario (QAS ready)', level=1)
files = sorted([f for f in glob.glob(os.path.join(payload_dir, '*.json'))])

table = doc.add_table(rows=1, cols=6)
headers = ['Archivo', 'KSCHL', 'Secuencia', 'Tipo', 'CODDESC', 'Referencia']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

for f in files:
    with open(f, encoding='utf-8') as fh:
        data = json.load(fh)
    kschl = str(data.get('coddescSap', ''))
    sec = str(data.get('indJerarq', ''))
    tipo = 'COMBO' if kschl in ('ZK96', 'ZR96') else 'SIMPLE/ESCALA/RECARGO'
    coddesc = str(data.get('coddesc', ''))
    ref = f"RegCond={data.get('regCond', '')} Promo={data.get('promo', '')}"
    row = table.add_row().cells
    row[0].text = os.path.basename(f)
    row[1].text = kschl
    row[2].text = sec
    row[3].text = tipo
    row[4].text = coddesc
    row[5].text = ref

doc.add_paragraph('Ubicacion de referencia de payloads: C:/Users/yejc2/source/repos/ROAD_TOLEDANO/ROADWedAPI/payloads/sap_descuentos/')

doc.add_heading('11. Notas de compatibilidad BOF/RDC7/HH', level=1)
notes = [
    'Compatibilidad historica: CODDESC NULL se interpreto como simple/escalado en implementaciones legacy.',
    'Decision actual del proyecto: persistir CODDESC cuando SAP lo envie para mejorar trazabilidad y correlacion funcional.',
    'Para combos: CODDESC + PTIPO=6 y detalle en P_DESCUENTO_COMBO_DET.'
]
for n in notes:
    p = doc.add_paragraph(n)
    p.style = 'List Bullet'

doc.save(doc_path)
print('UPDATED', doc_path)
print('BACKUP', backup_path)
